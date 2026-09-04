"""book formatter to publish in KDP."""

# Standard library imports.
from operator import index

import manuscript

# Third-party imports
from PySide6.QtWidgets import (
    QApplication,
    QWidget,
    QHBoxLayout,
    QVBoxLayout,
    QLabel,
    QComboBox,
    QPushButton,
    QFileDialog,
    QTextEdit,
    QProgressBar,
    QMenuBar,
    QDialog,
    QMessageBox,
)
from PySide6.QtCore import Qt
from PySide6.QtGui import (
    QImage,
    QTextBlockFormat,
    QTextCharFormat,
    QTextCursor,
)
from docx import Document

# Local imports
import constants
import book_layout
import theme
import front_matter

# Classes
class PagePreview(QTextEdit):
    def wheelEvent(self, event) -> None:
        event.ignore()

class bookformwindow(QWidget):
    """Controls the BookForm application window."""

    def __init__(self) -> None:
        super().__init__()

        # application state
        self.formatting_pending = False
        self.current_page = 0
        self.pages = []
        self.paragraphs = []
        self.page_starts_with_continuation = []
        self.gutter_width = 0.375
        self.page_types = []

        self.front_matter = front_matter.FrontMatter()
        
        # window configuration
        self.setWindowTitle("BookForm")
        self.resize(800, 600)
        self.setStyleSheet(theme.BOOKFORM_STYLESHEET)

        # layout setup
        self.outer_layout = QVBoxLayout(self)
        self.outer_layout.setContentsMargins(0, 0, 0, 0)
        self.outer_layout.setSpacing(0)

        self.layout = QHBoxLayout()

        self.menu_bar = QMenuBar()

        self.menu_bar.setObjectName("mainMenuBar")

        self.file_menu = self.menu_bar.addMenu(
            "File"
        )

        self.load_manuscript_action = (
            self.file_menu.addAction(
                "Load Manuscript..."
            )
        )

        self.file_menu.addSeparator()

        self.front_matter_menu = (
            self.file_menu.addMenu(
                "Front Matter"
            )       
        )

        self.title_page_action = (
            self.front_matter_menu.addAction(
                "Title Page"
            )
        )

        self.title_page_action.triggered.connect(
            self.configure_title_page
        )

        self.copyright_action = (
            self.front_matter_menu.addAction(
                "Copyright"
            )
        )

        self.dedication_action = (
            self.front_matter_menu.addAction(
                "Dedication"
            )
        )

        self.dedication_action.triggered.connect(
            self.configure_dedication
        )

        self.map_action = (
            self.front_matter_menu.addAction(
                "Map"
            )
        )

        self.trigger_warnings_action = (
            self.front_matter_menu.addAction(
                "Trigger Warnings"
            )
        )

        self.load_manuscript_action.triggered.connect(
            self.load_manuscript
        )

        self.outer_layout.addWidget(
            self.menu_bar
        )

        self.outer_layout.addLayout(
            self.layout
        )

        # panel creation
        self.controls_panel = QWidget()
        self.preview_panel = QWidget()
        self.layout.addWidget(self.controls_panel)
        self.layout.addWidget(self.preview_panel)
        self.controls_layout = QVBoxLayout()
        self.controls_panel.setLayout(self.controls_layout)
        self.preview_layout = QVBoxLayout()
        self.preview_panel.setLayout(self.preview_layout)

        # formatting controls
        trim_label = QLabel("Trim Size")

        self.trim_combo = QComboBox()
        self.trim_combo.addItem("6 by 9 inches")

        self.controls_layout.addWidget(trim_label)
        self.controls_layout.addWidget(self.trim_combo)

        font_label = QLabel("Body Font")

        self.font_combo = QComboBox()
        self.font_combo.addItems([
            "Garamond",
            "Times New Roman",
            "Georgia",
            "Arial",
            "Palatino Linotype",
        ])

        self.font_combo.activated.connect(
            self.mark_formatting_pending
        )

        self.controls_layout.addWidget(font_label)
        self.controls_layout.addWidget(self.font_combo)

        font_size_label = QLabel("Font Size")

        self.font_size_combo = QComboBox()
        self.font_size_combo.addItems([
            "8.5 pt",
            "9 pt",
            "10 pt",
            "11 pt",
            "12 pt",
        ])

        self.font_size_combo.activated.connect(
            self.mark_formatting_pending
        )

        self.controls_layout.addWidget(font_size_label)
        self.controls_layout.addWidget(self.font_size_combo)

        line_spacing_label = QLabel("Line Spacing")

        self.line_spacing_combo = QComboBox()
        self.line_spacing_combo.addItems([
            "1.0",
            "1.15",
            "1.5",
        ])

        self.line_spacing_combo.activated.connect(
            self.mark_formatting_pending
        )

        self.controls_layout.addWidget(line_spacing_label)
        self.controls_layout.addWidget(self.line_spacing_combo)

        margin_profile_label = QLabel("Margin Profile")

        self.margin_profile_combo = QComboBox()
        self.margin_profile_combo.addItems(
            constants.MARGIN_PROFILES.keys()
        )

        self.margin_profile_combo.setCurrentText(
            constants.DEFAULT_MARGIN_PROFILE
        )

        self.margin_profile_combo.activated.connect(
            self.mark_formatting_pending
        )

        self.controls_layout.addWidget(margin_profile_label)
        self.controls_layout.addWidget(self.margin_profile_combo)

        self.apply_changes_button = QPushButton("Apply Changes")

        self.formatting_status_label = QLabel("")
        self.controls_layout.addWidget(
            self.formatting_status_label
        )

        self.formatting_progress = QProgressBar()
        self.formatting_progress.setRange(0, 0)
        self.formatting_progress.hide()

        self.controls_layout.addWidget(
            self.formatting_progress
        )

        self.apply_changes_button.clicked.connect(
            self.apply_changes
        )

        self.controls_layout.addWidget(
            self.apply_changes_button
        )

        # manuscript controls
        self.controls_layout.addStretch()

        # preview setup
        self.preview_label = QLabel("Book Preview")

        self.preview_label.setAlignment(
            Qt.AlignmentFlag.AlignCenter
        )

        self.preview_layout.addWidget(
            self.preview_label
        )

        # future two-page spread
        self.left_page_preview = self.create_page_preview()
        self.right_page_preview = self.create_page_preview()

        spread_layout = QHBoxLayout()
        spread_layout.addWidget(self.left_page_preview)
        spread_layout.addWidget(self.right_page_preview)

        # existing preview used by current pagination
        self.page_preview = self.create_page_preview()

        self.measurement_device = QImage(
            int(constants.PRINT_PAGE_WIDTH_POINTS),
            int(constants.PRINT_PAGE_HEIGHT_POINTS),
            QImage.Format.Format_ARGB32,
        )

        self.measurement_device.setDotsPerMeterX(
            constants.PRINT_DOTS_PER_METER
        )

        self.measurement_device.setDotsPerMeterY(
            constants.PRINT_DOTS_PER_METER
        )

        self.page_preview.document().documentLayout().setPaintDevice(
            self.measurement_device
        )
        self.update_measurement_area()

        self.preview_layout.addLayout(spread_layout)

        self.kdp_warning_label = QLabel("")
        self.kdp_warning_label.setWordWrap(True)
        self.preview_layout.addWidget(self.kdp_warning_label)     
        
        # page navigation
        navigation_layout = QHBoxLayout()
        
        self.previous_button = QPushButton("Previous")
        self.next_button = QPushButton("Next")
        self.page_number_label = QLabel("Page 0 of 0")
        
        self.previous_button.clicked.connect(self.previous_page)
        self.next_button.clicked.connect(self.next_page)
                
        navigation_layout.addWidget(self.previous_button)
        navigation_layout.addWidget(self.page_number_label)
        navigation_layout.addWidget(self.next_button)
        
        self.preview_layout.addLayout(navigation_layout)

    # Methods
    def load_manuscript(self) -> None:
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Select Manuscript",
                "",
                "Word Documents (*.docx)"
            )
    
            if file_path:
                document = Document(file_path)
                styles_found = set()
    
                self.paragraphs = []
    
                for paragraph in document.paragraphs:
                    styles_found.add(paragraph.style.name)
    
                    if paragraph.text.strip():
                        self.paragraphs.append(paragraph.text)

                
                self.pages = []
                self.current_page = 0
                self.formatting_pending = True
                self.formatting_status_label.setText(
                    "Formatting changes pending"
                )

                self.left_page_preview.clear()
                self.right_page_preview.clear()

                self.page_number_label.setText("Page 0 of 0")

    def configure_title_page(self) -> None:
        if not self.paragraphs:
            QMessageBox.warning(
                self,
                "No Manuscript Loaded",
                "Load a manuscript before identifying the title page.",
            )
            return

        dialog = front_matter.TitlePageDialog(self)

        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        title, subtitle, author = dialog.values()

        title_page = front_matter.find_title_page(
            self.paragraphs,
            title,
            subtitle,
            author,
        )

        if title_page is None:
            QMessageBox.warning(
                self,
                "Title Page Not Found",
                "BookForm could not find all of the specified "
                "title page text in the manuscript.",
            )
            return

        self.front_matter.title_page = title_page

        self.formatting_pending = True
        self.formatting_status_label.setText(
            "Formatting changes pending"
        )

        print(
            "Title page range:",
            title_page.start_index,
            "to",
            title_page.end_index,
        )

        print("Title page paragraphs:")

        for index in range(
            title_page.start_index,
            title_page.end_index + 1,
        ):
            print(
                index,
                repr(self.paragraphs[index]),
            )

        QMessageBox.information(
            self,
            "Title Page Found",
            "The title page was found in the manuscript.",
        )

    def configure_dedication(self) -> None:
        if not self.paragraphs:
            QMessageBox.warning(
                self,
                "No Manuscript Loaded",
                "Load a manuscript before identifying the dedication.",
            )
            return
        
        dialog = front_matter.DedicationDialog(
            self
        )

        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        dedication_text = dialog.value()

        if not dedication_text:
            QMessageBox.warning(
                self,
                "No Dedication",
                "No dedication text was provided.",
            )
            return

        dedication = front_matter.find_text_section(
            self.paragraphs,
            dedication_text,
        )  

        if dedication is None:
            QMessageBox.warning(
                self,
                "Dedication Not Found",
                "BookForm could not find that "
                "dedication text in the manuscript.",
            )
            return

        self.front_matter.dedication = dedication

        self.formatting_pending = True
        self.formatting_status_label.setText(
            "Formatting changes pending"
        )

        QMessageBox.information(
            self,
            "Dedication Ready",
            "The dedication was found in the manuscript.",
        )

    def mark_formatting_pending(self, _index: int) -> None:
        self.formatting_pending = True

        self.formatting_status_label.setText(
            "Formatting changes pending"
        )

    def create_page_preview(self) -> QTextEdit:
        preview = PagePreview()

        preview.document().setDocumentMargin(0)

        preview.setFixedSize(
            constants.PREVIEW_WIDTH_PIXELS,
            constants.PREVIEW_HEIGHT_PIXELS,
        )

        preview.setReadOnly(True)
        preview.setFocusPolicy(Qt.NoFocus)
        preview.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        preview.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        return preview

    
    def apply_page_margins(self) -> None:
        profile_name = self.margin_profile_combo.currentText()
        profile = constants.MARGIN_PROFILES[profile_name]

        top_margin = book_layout.inches_to_pixels(profile["top"])
        bottom_margin = book_layout.inches_to_pixels(profile["bottom"])
        outside_margin = book_layout.inches_to_pixels(profile["outside"])
        gutter_margin = book_layout.inches_to_pixels(self.gutter_width)

        page_number = self.current_page + 1

        if page_number % 2 == 1:
            left_margin = gutter_margin
            right_margin = outside_margin
        else:
            left_margin = outside_margin
            right_margin = gutter_margin

        self.left_page_preview.setStyleSheet(f"""
            background-color: white;
            border: 0px solid gray;
            padding-top: {top_margin}px;
            padding-bottom: {bottom_margin}px;
            padding-left: {outside_margin}px;
            padding-right: {gutter_margin}px;
        """)

        self.right_page_preview.setStyleSheet(f"""
            background-color: white;
            border: 0px solid gray;
            padding-top: {top_margin}px;
            padding-bottom: {bottom_margin}px;
            padding-left: {gutter_margin}px;
            padding-right: {outside_margin}px;
        """)

    def apply_font_settings(
        self,
        preview: QTextEdit,
    ) -> None:
        font_name = self.font_combo.currentText()
        size_text = self.font_size_combo.currentText()
        font_size = float(size_text.split()[0])

        if preview is self.page_preview:
            applied_font_size = font_size
        else:
            applied_font_size = (
                font_size
                * constants.PIXELS_PER_INCH
                / preview.logicalDpiY()
            )

        font = preview.font()
        font.setFamily(font_name)
        font.setPointSizeF(applied_font_size)

        preview.setFont(font)
        preview.document().setDefaultFont(font)

        cursor = QTextCursor(
            preview.document()
        )
        cursor.select(
            QTextCursor.SelectionType.Document
        )

        body_format = QTextCharFormat()
        body_format.setFont(font)

        cursor.setCharFormat(
            body_format
        )

    def show_page(self) -> None:
        if not self.pages:
            return

        self.apply_page_margins()

        right_index = self.current_page
        left_index = right_index - 1

        if left_index >= 0:
            self.left_page_preview.setPlainText(
                self.pages[left_index]
            )
        else:
            self.left_page_preview.clear()

        if right_index < len(self.pages):
            self.right_page_preview.setPlainText(
                self.pages[right_index]
            )
        else:
            self.right_page_preview.clear()

        left_continues = (
            left_index >= 0
            and self.page_starts_with_continuation[left_index]
        )

        right_continues = (
            right_index < len(self.pages)
            and self.page_starts_with_continuation[right_index]
        )

        self.apply_font_settings(
            self.left_page_preview
        )
        self.apply_font_settings(
            self.right_page_preview
        )

        if left_index >= 0:
            left_page_type = self.page_types[left_index]

            if left_page_type == "title_page":
                self.apply_title_page_formatting(
                    self.left_page_preview
                )

            elif left_page_type == "dedication":
                self.apply_dedication_formatting(
                    self.left_page_preview
                )

            else:
                self.apply_paragraph_formatting(
                    self.left_page_preview,
                    left_continues,
                )

        if right_index < len(self.pages):
            right_page_type = self.page_types[right_index]

            if right_page_type == "title_page":
                self.apply_title_page_formatting(
                    self.right_page_preview
                )

            elif right_page_type == "dedication":
                self.apply_dedication_formatting(
                    self.right_page_preview
                )

            else:
                self.apply_paragraph_formatting(
                    self.right_page_preview,
                    right_continues,
                )

        if left_index >= 0 and right_index < len(self.pages):
            self.page_number_label.setText(
                f"Pages {left_index + 1}–{right_index + 1} "
                f"of {len(self.pages)}"
            )
        elif right_index < len(self.pages):
            self.page_number_label.setText(
            f"Page {right_index + 1} of {len(self.pages)}"
        )
        elif left_index >= 0:
            self.page_number_label.setText(
                f"Page {left_index + 1} of {len(self.pages)}"
            )


    def previous_page(self) -> None:
        if self.current_page >= 2:
            self.current_page -= 2
            self.show_page()


    def next_page(self) -> None:
        if self.current_page + 1 < len(self.pages):
            self.current_page += 2
            self.show_page()


    def update_pages(self) -> None:
        self.pages = []
        self.page_starts_with_continuation = []
        self.page_types = []

        current_page_text = ""
        current_page_starts_with_continuation = False

        for paragraph_index, paragraph in enumerate(
            self.paragraphs
        ):
            if paragraph_index % 25 == 0:
                QApplication.processEvents()

            title_page = self.front_matter.title_page

            if (
                title_page is not None
                and title_page.start_index is not None
                and title_page.end_index is not None
                and title_page.start_index
                <= paragraph_index
                <= title_page.end_index
            ):
                if paragraph_index == title_page.start_index:
                    if current_page_text:
                        self.add_page(
                            current_page_text,
                            current_page_starts_with_continuation,
                        )

                        current_page_text = ""
                        current_page_starts_with_continuation = False

                    title_page_text = "\n".join(
                        self.paragraphs[
                            title_page.start_index:
                            title_page.end_index + 1
                        ]
                    )

                    self.add_page(
                        title_page_text,
                        False,
                        page_type="title_page",
                    )

                continue

            dedication = self.front_matter.dedication

            if (
                dedication is not None
                and dedication.start_index is not None
                and dedication.end_index is not None
                and dedication.start_index
                <= paragraph_index
                <= dedication.end_index
            ):
                if paragraph_index == dedication.start_index:
                    if current_page_text:
                        self.add_page(
                            current_page_text,
                            current_page_starts_with_continuation,
                        )

                        current_page_text = ""
                        current_page_starts_with_continuation = False

                    dedication_text = "\n".join(
                        self.paragraphs[
                            dedication.start_index:
                            dedication.end_index + 1
                        ]
                    )

                    self.add_page(
                        dedication_text,
                        False,
                        page_type="dedication",
                    )

                continue

            if manuscript.is_chapter_heading(paragraph):
                if current_page_text:
                    self.add_page(
                        current_page_text,
                        current_page_starts_with_continuation,
                    )
                    current_page_text = ""
                    current_page_starts_with_continuation = False

            paragraph_remaining = paragraph
            paragraph_continues = False
        
            while paragraph_remaining:
                if current_page_text:
                    test_text = (
                        current_page_text
                        + "\n"
                        + paragraph_remaining
                    )
                    test_starts_with_continuation = (
                        current_page_starts_with_continuation
                    )
                else:
                    test_text = paragraph_remaining
                    test_starts_with_continuation = paragraph_continues
        
                if self.text_fits_page(
                    test_text,
                    test_starts_with_continuation,
                ):
                    if not current_page_text:
                        current_page_starts_with_continuation = (
                            paragraph_continues
                        )

                    current_page_text = test_text
                    paragraph_remaining = ""
                else:
                    if current_page_text:
                        fitting_text, paragraph_remaining = (
                            self.split_paragraph_to_fit(
                                paragraph_remaining,
                                current_page_text,
                                current_page_starts_with_continuation,
                            )
                        )

                        if fitting_text:
                            self.add_page(
                                current_page_text
                                + "\n"
                            + fitting_text,
                            current_page_starts_with_continuation,
                            )

                            if paragraph_remaining:
                                paragraph_continues = True

                        else:
                            self.add_page(
                                current_page_text,
                                current_page_starts_with_continuation,
                            )

                        current_page_text = ""
                        current_page_starts_with_continuation = False

                    else:
                        fitting_text, paragraph_remaining = (
                            self.split_paragraph_to_fit(
                                paragraph_remaining,
                                first_block_continues=paragraph_continues,
                            )
                        )

                        self.add_page(
                            fitting_text,
                            paragraph_continues,
                        )

                        if paragraph_remaining:
                            paragraph_continues = True
        
        if current_page_text:
            self.add_page(
                current_page_text,
                current_page_starts_with_continuation,
            )

        if self.pages:
            print(f"Page count: {len(self.pages)}")
            print(
                f"Word count: "
                f"{sum(len(paragraph.split()) for paragraph in self.paragraphs)}"
            )
            print(
                f"Measurement area: "
                f"{self.measurement_width} x "
                f"{self.measurement_height}"
            )
            page_count = len(self.pages)

            new_gutter_width = book_layout.calculate_gutter_width(page_count)

            if new_gutter_width != self.gutter_width:
                self.gutter_width = new_gutter_width
                self.update_measurement_area()
                self.update_pages()
                return

            if page_count > constants.KDP_MAX_PAGE_COUNT:
                pages_over = page_count - constants.KDP_MAX_PAGE_COUNT

                self.kdp_warning_label.setText(
                    f"{page_count} pages — "
                    f"{pages_over} pages over the KDP 6 × 9 limit."
                )
            else:
                self.kdp_warning_label.setText("")

    def update_measurement_area(self) -> None:
        profile_name = self.margin_profile_combo.currentText()
        profile = constants.MARGIN_PROFILES[profile_name]

        top_margin = book_layout.inches_to_points(
            profile["top"]
        )
        bottom_margin = book_layout.inches_to_points(
            profile["bottom"]
        )
        outside_margin = book_layout.inches_to_points(
            profile["outside"]
        )
        gutter_margin = book_layout.inches_to_points(
            self.gutter_width
        )

        self.measurement_width = (
            constants.PRINT_PAGE_WIDTH_POINTS
            - gutter_margin
            - outside_margin
        )

        self.measurement_height = (
            constants.PRINT_PAGE_HEIGHT_POINTS
            - top_margin
            - bottom_margin
        )

    def text_fits_page(
        self,
        text: str,
        first_block_continues: bool = False,
    ) -> bool:
        self.page_preview.setPlainText(text)

        self.apply_font_settings(self.page_preview)
        self.apply_paragraph_formatting(
            self.page_preview,
            first_block_continues,
        )

        self.page_preview.document().setTextWidth(
            self.measurement_width
        )

        document_height = self.page_preview.document().size().height()

        return document_height <= self.measurement_height


    def split_paragraph_to_fit(
        self,
        paragraph: str,
        current_page_text: str = "",
        first_block_continues: bool = False,
    ) -> tuple[str, str]:
        words = paragraph.split()
        fitting_words = []

        for index, word in enumerate(words):
            candidate = " ".join(fitting_words + [word])

            if current_page_text:
                test_text = (
                    current_page_text
                    + "\n"
                    + candidate
                )
            else:
                test_text = candidate

            if self.text_fits_page(
                test_text,
                first_block_continues,
            ):
                fitting_words.append(word)
            else:
                if not fitting_words:
                    if current_page_text:
                        return "", paragraph

                    return word, " ".join(words[index + 1:])

                return (
                    " ".join(fitting_words),
                    " ".join(words[index:])
                )

        return " ".join(fitting_words), ""

    def update_font_size(self) -> None:
        self.apply_font_settings(self.page_preview)

        self.update_pages()
        self.show_page()

    def update_font(self) -> None:
        self.apply_font_settings(self.page_preview)

        self.update_pages()
        self.show_page()

    def apply_paragraph_formatting(
        self,
        preview: QTextEdit,
        first_block_continues: bool = False,
    ) -> None:
        spacing = float(self.line_spacing_combo.currentText())
        line_height = spacing * 100

        if preview is self.page_preview:
            first_line_indent = book_layout.inches_to_points(
                constants.FIRST_LINE_INDENT_INCHES
            )
            heading_space_after = book_layout.inches_to_points(
                constants.CHAPTER_HEADING_SPACE_AFTER_INCHES
            )
        else:
            first_line_indent = book_layout.inches_to_pixels(
                constants.FIRST_LINE_INDENT_INCHES
            )
            heading_space_after = book_layout.inches_to_pixels(
                constants.CHAPTER_HEADING_SPACE_AFTER_INCHES
            )

        document = preview.document()
        block = document.begin()

        previous_was_heading = False
        block_number = 0

        while block.isValid():
            text = block.text().strip()
            is_heading = manuscript.is_chapter_heading(text)

            block_format = QTextBlockFormat()

            block_format.setLineHeight(
                line_height,
                QTextBlockFormat.LineHeightTypes.ProportionalHeight.value
            )

            if is_heading:
                block_format.setTextIndent(0)
                block_format.setAlignment(
                    Qt.AlignmentFlag.AlignCenter
                )
                block_format.setBottomMargin(
                    heading_space_after
                )

            elif previous_was_heading:
                block_format.setTextIndent(0)

            elif block_number == 0 and first_block_continues:
                block_format.setTextIndent(0)

            else:
                block_format.setTextIndent(
                    first_line_indent
                )

            cursor = QTextCursor(block)
            cursor.setBlockFormat(block_format)

            if is_heading:
                heading_format = QTextCharFormat()
                if preview is self.page_preview:
                    heading_font_size = constants.CHAPTER_HEADING_FONT_SIZE
                else:
                    heading_font_size = (
                        constants.CHAPTER_HEADING_FONT_SIZE
                        * constants.PIXELS_PER_INCH
                        / preview.logicalDpiY()
                    )

                heading_format.setFontPointSize(
                    heading_font_size
                )

                cursor.movePosition(
                    QTextCursor.MoveOperation.StartOfBlock
                )
                cursor.movePosition(
                    QTextCursor.MoveOperation.EndOfBlock,
                    QTextCursor.MoveMode.KeepAnchor,
                )

                cursor.setCharFormat(heading_format)

            previous_was_heading = is_heading
            block_number += 1
            block = block.next()

    def apply_title_page_formatting(
        self,
        preview: QTextEdit,
    ) -> None:
        document = preview.document()
        block = document.begin()
        block_number = 0

        while block.isValid():
            block_format = QTextBlockFormat()

            block_format.setAlignment(
                Qt.AlignmentFlag.AlignCenter
            )
            block_format.setTextIndent(0)

            if block_number == 0:
                if preview is self.page_preview:
                    top_space = book_layout.inches_to_points(
                        constants.TITLE_PAGE_TOP_SPACE_INCHES
                    )
                else:
                    top_space = book_layout.inches_to_pixels(
                        constants.TITLE_PAGE_TOP_SPACE_INCHES
                    )

                block_format.setTopMargin(
                    top_space
                )

            if block_number == 1:
                if preview is self.page_preview:
                    author_space = book_layout.inches_to_points(
                        constants.TITLE_PAGE_AUTHOR_SPACE_INCHES
                    )
                else:
                    author_space = book_layout.inches_to_pixels(
                        constants.TITLE_PAGE_AUTHOR_SPACE_INCHES
                    )

                block_format.setBottomMargin(
                    author_space
                )

            cursor = QTextCursor(block)
            cursor.setBlockFormat(block_format)

            text_format = QTextCharFormat()

            if block_number == 0:
                font_size = 20
            elif block_number == 1:
                font_size = 14
            else:
                font_size = 12

            if preview is not self.page_preview:
                font_size = (
                    font_size
                    * constants.PIXELS_PER_INCH
                    / preview.logicalDpiY()
                )

            text_format.setFontPointSize(
                font_size
            )

            cursor.movePosition(
                QTextCursor.MoveOperation.StartOfBlock
            )
            cursor.movePosition(
                QTextCursor.MoveOperation.EndOfBlock,
                QTextCursor.MoveMode.KeepAnchor,
            )

            cursor.setCharFormat(
                text_format
            )

            block_number += 1
            block = block.next()

    def apply_dedication_formatting(
        self,
        preview: QTextEdit,
    ) -> None:
        if preview is self.page_preview:
            top_space = book_layout.inches_to_points(
                constants.DEDICATION_TOP_SPACE_INCHES
            )
        else:
            top_space = book_layout.inches_to_pixels(
                constants.DEDICATION_TOP_SPACE_INCHES
            )

        document = preview.document()
        block = document.begin()
        block_number = 0

        while block.isValid():
            block_format = QTextBlockFormat()

            block_format.setAlignment(
                Qt.AlignmentFlag.AlignCenter
            )
            block_format.setTextIndent(0)

            if block_number == 0:
                block_format.setTopMargin(
                    top_space
                )

            cursor = QTextCursor(block)
            cursor.setBlockFormat(
                block_format
            )

            block_number += 1
            block = block.next()

    def line_spacing_changed(self) -> None:
        self.apply_paragraph_formatting(self.page_preview)
        self.update_pages()
        self.show_page()

    def margin_profile_changed(self) -> None:
        self.update_measurement_area()
        self.update_pages()
        self.show_page()

    def add_page(
        self,
        text: str,
        starts_with_continuation: bool = False,
        page_type: str = "body",
    ) -> None:
        self.pages.append(text)

        self.page_starts_with_continuation.append(
            starts_with_continuation
        )

        self.page_types.append(
            page_type
        )

    def apply_changes(self) -> None:
        if not self.paragraphs:
            return

        self.formatting_status_label.setText(
            "Formatting manuscript, please wait..."
        )
        self.formatting_progress.show()
        self.apply_changes_button.setEnabled(False)

        QApplication.processEvents()

        try:
            self.apply_font_settings(
                self.page_preview
            )

            self.update_measurement_area()
            self.update_pages()

            self.current_page = 0
            self.show_page()

            self.formatting_pending = False
            self.formatting_status_label.setText("")

        finally:
            self.formatting_progress.hide()
            self.apply_changes_button.setEnabled(True)


# Functions

def is_chapter_heading(text: str) -> bool:
    return bool(re.match(r"^\d+\.\s+\S", text.strip()))

def has_numbering(paragraph) -> bool:
    paragraph_properties = paragraph._p.pPr

    return (
        paragraph_properties is not None
        and paragraph_properties.numPr is not None
    )


def main():
    app = QApplication([])

    window = bookformwindow()
    window.show()

    app.exec()

if __name__ == "__main__":
    main()